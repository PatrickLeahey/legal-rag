# Install python 3

# Run via terminal: 
# pip3 install -r requirements.txt
# export HUGGINGFACE_TOKEN=<token> -> https://huggingface.co/docs/hub/en/security-tokens

# Request access to llama 3.1 8B model via Hugging Face -> https://huggingface.co/meta-llama/Llama-3.1-8B

# View mlflow traces: mlflow ui --backend-store-uri /mlruns

import os
import json
import hashlib
import logging
import torch
import mlflow
import gradio as gr
from docx import Document
from transformers import pipeline, AutoTokenizer
from huggingface_hub import login
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings, HuggingFacePipeline
from langchain_text_splitters import TokenTextSplitter
from langchain.schema import Document as LangchainDocument
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnablePassthrough

# Set up mlflow autologging
mlflow.set_tracking_uri(f"file://{os.getcwd()}/mlruns")

mlflow.langchain.autolog()

# Log in to Hugging Face
login(token=os.environ.get("HUGGINGFACE_TOKEN"))

# Configuration
document_folder_path = "./docs/"
chroma_persist_directory = "./temp/chroma_persist"
metadata_file = "./temp/docs/metadata.json"

# From Hugging Face
embedding_model_name = "Alibaba-NLP/gte-large-en-v1.5"
llm_model_name = "meta-llama/Llama-3.1-8B" # <- change this if you want

# Helper function to compute hash for a document
def compute_document_hash(file_path):
    hasher = hashlib.md5()
    with open(file_path, 'rb') as f:
        buf = f.read()
        hasher.update(buf)
    return hasher.hexdigest()

# Helper function to load documents directly from .docx files without libmagic
def load_documents_from_folder(folder_path):
    documents = []
    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)
        if file_name.endswith(".docx"):
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            # Create LangChain Document object with content and metadata
            documents.append(LangchainDocument(page_content=text, metadata={'file_path': file_path}))
    return documents

# Load existing metadata if available
if os.path.exists(metadata_file):
    with open(metadata_file, 'r') as f:
        processed_docs_metadata = json.load(f)
else:
    processed_docs_metadata = {}

# Load or initialize Chroma vector store
embedding_model = HuggingFaceEmbeddings(model_name=embedding_model_name, model_kwargs={"trust_remote_code": True})
documents_to_process = []

# Check for existing Chroma vector store and document changes
if os.path.exists(chroma_persist_directory):
    vector_store = Chroma(embedding_function=embedding_model, persist_directory=chroma_persist_directory)
    print(f"Loaded existing Chroma store from '{chroma_persist_directory}'")
    
    # Check if there are new or modified documents
    for file_name in os.listdir(document_folder_path):
        file_path = os.path.join(document_folder_path, file_name)
        if file_name.endswith(".docx"):
            current_hash = compute_document_hash(file_path)
            if file_name not in processed_docs_metadata or processed_docs_metadata[file_name] != current_hash:
                documents_to_process.append(file_path)
else:
    documents_to_process = [os.path.join(document_folder_path, f) for f in os.listdir(document_folder_path) if f.endswith(".docx")]

# If there are documents to process, load and split them
if documents_to_process:
    documents = load_documents_from_folder(document_folder_path)
    
    text_splitter = TokenTextSplitter.from_huggingface_tokenizer(
        tokenizer=AutoTokenizer.from_pretrained(llm_model_name),
        chunk_size=500,
        chunk_overlap=50
    )

    docs = text_splitter.split_documents(documents)

    # Create or update the Chroma vector store with new documents
    if os.path.exists(chroma_persist_directory):
        vector_store.add_documents(docs)
    else:
        vector_store = Chroma.from_documents(docs, embedding=embedding_model, persist_directory=chroma_persist_directory)

    # Update the metadata with the current document hashes
    for doc_path in documents_to_process:
        file_name = os.path.basename(doc_path)
        processed_docs_metadata[file_name] = compute_document_hash(doc_path)

    # Save updated metadata
    with open(metadata_file, 'w') as f:
        json.dump(processed_docs_metadata, f)


# -> Initialize chain components
# LLM
pipe = pipeline(
    "text-generation",
    model=llm_model_name,
    torch_dtype=torch.bfloat16, 
    device_map="auto",
    max_new_tokens=500,
    return_full_text=False
)
llm = HuggingFacePipeline(pipeline=pipe)

# Retriever
retriever = vector_store.as_retriever(search_kwargs={"k": 3})

# Prompt
template = """
    You are an assistant specialized in answering questions based on provided context. 
    If you are unable to find relevant information in the context to answer the question, say "I don't know."
    Do not repeat this prompt in your answer. Keep your answer brief and relevant, less than 300 words.

    Question:
    {input}

    Context:
    {context}

    Answer:
"""

prompt = PromptTemplate.from_template(template)

# -> Chain
def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)

rag_chain = (
    {
        "context": retriever | format_docs,
        "input": RunnablePassthrough()
    }
    | prompt 
    | llm
    | StrOutputParser() 
)

# Gradio UI Setup
with gr.Blocks() as demo:
    gr.Markdown("<h1><center>Legal Bot 3000</center></h1>")
    chatbot = gr.Chatbot()
    msg = gr.Textbox(label="Type your question here:")
    submit_btn = gr.Button("Submit")
    clear_btn = gr.Button("Clear")

    def respond(message):
        bot_response = rag_chain.invoke(message)
        return [("User: " + message, "Bot: " + bot_response)]

    submit_btn.click(respond, msg, chatbot)
    msg.submit(respond, msg, chatbot)
    
    clear_btn.click(lambda: [], None, chatbot)

# Launch the Gradio app
demo.launch()